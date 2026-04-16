import hashlib
import os
from typing import List, Tuple

from utils.pdf_reader import PDFReader


def compute_content_hash(text: str) -> str:
    """计算文本内容的 SHA-256 哈希值，用于去重"""
    normalized = text.strip().lower()
    return hashlib.sha256(normalized.encode('utf-8')).hexdigest()


def scan_all_files(folder_path: str) -> List[Tuple[str, str]]:
    """
    扫描文件夹中所有支持的文件类型（pdf, docx, md）

    Returns:
        (file_path, file_type) 元组列表，file_type 为 'pdf'/'docx'/'md'
    """
    supported = {'.pdf': 'pdf', '.docx': 'docx', '.md': 'md'}
    results = []
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            ext = os.path.splitext(file)[1].lower()
            if ext in supported:
                results.append((os.path.join(root, file), supported[ext]))
    return results


class TextExtractor:
    """统一文本提取器，支持 PDF、Word、Markdown 三种格式"""

    def __init__(self):
        self.pdf_reader = PDFReader()

    def extract(self, file_path: str) -> Tuple[str, str]:
        """
        从文件中提取文本

        Returns:
            (text_content, file_type) 元组

        Raises:
            ValueError: 不支持的文件类型
        """
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.pdf':
            return self._extract_pdf(file_path), 'pdf'
        elif ext == '.docx':
            return self._extract_docx(file_path), 'docx'
        elif ext == '.md':
            return self._extract_md(file_path), 'md'
        else:
            raise ValueError(f"不支持的文件类型: {ext}")

    def _extract_pdf(self, file_path: str) -> str:
        """使用现有 PDFReader 提取 PDF 文本"""
        return self.pdf_reader.extract_text(file_path)

    def _extract_docx(self, file_path: str) -> str:
        """提取 Word 文档文本（段落 + 表格）"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")

        doc = Document(file_path)
        text_parts = []

        # 提取段落
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # 提取表格
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(' | '.join(row_text))

        text = '\n'.join(text_parts)
        if not text.strip():
            raise ValueError("Word 文档中未提取到任何文本内容")
        return text

    def _extract_md(self, file_path: str) -> str:
        """读取 Markdown 文件内容"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()

        if not text.strip():
            raise ValueError("Markdown 文件内容为空")
        return text
